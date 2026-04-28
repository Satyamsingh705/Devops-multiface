from docx import Document
from docx.shared import Pt

md_path = "report/CAPSTONE_FULL.md"
docx_path = "report/CAPSTONE_REPORT.docx"

def add_paragraph_with_style(doc, text, bold=False, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def main():
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    with open(md_path, 'r', encoding='utf-8') as f:
        data = f.read()

    pages = data.split('<!-- PAGE BREAK -->')
    for i, page in enumerate(pages, start=1):
        lines = [l.rstrip() for l in page.strip().split('\n') if l.strip()]
        for line in lines:
            if line.startswith('**') and line.endswith('**'):
                add_paragraph_with_style(doc, line.strip('*'), bold=True, size=16)
            elif line.startswith('```'):
                # simple code block handling
                add_paragraph_with_style(doc, line, size=10)
            else:
                add_paragraph_with_style(doc, line, size=12)
        if i != len(pages):
            doc.add_page_break()

    doc.save(docx_path)
    print(f"Saved {docx_path}")

if __name__ == '__main__':
    main()
